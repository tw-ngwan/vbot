from docx import Document

# Creates the word document, assuming you already have the report

import os 
# os.chdir('../../Oxford')

def add_bold_underline(text, target):
    doc = Document()

    parts = text.split("\n")
    for part in parts:
        part = part.strip()
        if part.startswith(target):
            # Split the part into 'Trip Number' and 'Further text'
            # print(part)
            _, trip_number = part.split(target, 1)
            
            # Add 'Trip Number' with bold and underline formatting
            p = doc.add_paragraph()
            run = p.add_run(f"{target.strip()} {trip_number.strip()}")
            run.bold = True
            run.underline = True

        else:
            # Add other parts directly in a new paragraph
            doc.add_paragraph(part.strip())

    return doc

# # Example usage
# text = """ 1
# Further text"""
# target = "Trip Number"

# doc = add_bold_underline(text, target)
# doc.save("formatted_text.docx")


with open('report.txt', 'r') as f:
    text = f.read() 

target = "Trip Number:"

doc = add_bold_underline(text, target)
doc.save("trip-report.docx")
